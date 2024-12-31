import streamlit as st
from langchain_groq import ChatGroq
from langchain.document_loaders import WebBaseLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from bs4 import BeautifulSoup
import requests
from docx import Document
from PyPDF2 import PdfReader
from sentence_transformers import SentenceTransformer
import os
import json
from typing import Dict, List, Optional
from dotenv import load_dotenv
import numpy as np
import pandas as pd

# Load environment variables
load_dotenv()
GROQ_API_KEY = os.getenv('GROQ_API_KEY')

class ResumeTailor:
    def __init__(self):
        """Initialize the ResumeTailor with necessary components."""
        if not GROQ_API_KEY:
            raise ValueError("GROQ_API_KEY not found in environment variables")
            
        self.llm = ChatGroq(
            model="llama-3.1-70b-versatile",
            groq_api_key=GROQ_API_KEY,
            temperature=0,
            max_tokens=None,
            timeout=None,
            max_retries=2
        )
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text content from a PDF file using PyPDF2."""
        pdf_reader = PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text()
        return text
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text content from a DOCX file."""
        doc = Document(docx_file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    
    def parse_job_description(self, job_text: str, is_url: bool = True) -> Dict:
        """Extract job details from the provided URL or text."""
        if is_url:
            loader = WebBaseLoader(job_text)
            data = loader.load()
            job_content = data[0].page_content
        else:
            job_content = job_text
        
        # Use LLM to extract structured information
        prompt = f"""Please analyze the following job description and extract key information in JSON format.
        Return ONLY a JSON object with the following structure:
        {{
            "skills": ["skill1", "skill2", ...],
            "experience": "experience requirements",
            "responsibilities": ["responsibility1", "responsibility2", ...],
            "company": "company name",
            "title": "job title"
        }}

        Job Description:
        {job_content}
        """
        
        response = self.llm.invoke(prompt)
        try:
            # Extract content and clean it to ensure it's valid JSON
            content = str(response.content).strip()
            if content.startswith("```json"):
                content = content.split("```json")[1]
            if content.endswith("```"):
                content = content.rsplit("```", 1)[0]
            return json.loads(content.strip())
        except json.JSONDecodeError:
            # Fallback structure if parsing fails
            return {
                "skills": [],
                "experience": "Not specified",
                "responsibilities": [],
                "company": "Not specified",
                "title": "Not specified"
            }
    
    def match_skills(self, resume_text: str, job_requirements: Dict) -> Dict:
        """Match resume skills with job requirements using semantic similarity."""
        try:
            # Ensure skills is a list and not empty
            if not job_requirements.get('skills') or not isinstance(job_requirements['skills'], list):
                return {
                    'matched_skills': [],
                    'missing_skills': []
                }
            
            # Common tech variations dictionary
            tech_variations = {
                'ml': ['machine learning', 'ml'],
                'ai': ['artificial intelligence', 'ai'],
                'llm': ['large language model', 'llm', 'llama', 'gpt', 'language model'],
                'nlp': ['natural language processing', 'nlp'],
                'js': ['javascript', 'js'],
                'ts': ['typescript', 'ts'],
                'py': ['python', 'py'],
                'react': ['reactjs', 'react.js', 'react'],
                'node': ['nodejs', 'node.js', 'node'],
                'db': ['database', 'db'],
                'ui': ['user interface', 'ui'],
                'ux': ['user experience', 'ux'],
                'api': ['apis', 'api', 'restful', 'rest'],
                'aws': ['amazon web services', 'aws'],
                'gcp': ['google cloud platform', 'gcp'],
                'azure': ['microsoft azure', 'azure'],
                'k8s': ['kubernetes', 'k8s'],
                'ci/cd': ['continuous integration', 'continuous deployment', 'ci/cd', 'cicd'],
                'oop': ['object oriented programming', 'object-oriented', 'oop'],
                'cv': ['computer vision', 'cv']
            }
            
            # Preprocess resume text
            resume_text_lower = resume_text.lower()
            resume_sentences = [sent.strip() for sent in resume_text.split('.')]
            
            # First pass: Direct keyword matching with variations
            matched_skills = []
            remaining_skills = []
            
            for skill in job_requirements['skills']:
                skill_lower = skill.lower()
                # Check if this skill has known variations
                variations = []
                for var_key, var_list in tech_variations.items():
                    if skill_lower in var_list:
                        variations.extend(var_list)
                    # Also check if any variation is in the skill name
                    for var in var_list:
                        if var in skill_lower:
                            variations.extend(var_list)
                
                # Add common text variations
                variations.extend([
                    skill_lower,
                    skill_lower.replace(' ', ''),
                    skill_lower.replace('-', ''),
                    skill_lower.replace('.', ''),
                    skill_lower.replace('/', '')
                ])
                
                # Remove duplicates and empty strings
                variations = list(set(filter(None, variations)))
                
                if any(var in resume_text_lower for var in variations):
                    matched_skills.append(skill)
                else:
                    remaining_skills.append(skill)
            
            # Second pass: Semantic matching for remaining skills
            if remaining_skills:
                # Prepare embeddings for remaining skills and their variations
                skill_texts = []
                skill_map = {}  # Map expanded texts back to original skills
                
                for skill in remaining_skills:
                    skill_lower = skill.lower()
                    variations = []
                    # Add known variations
                    for var_list in tech_variations.values():
                        if any(var in skill_lower for var in var_list):
                            variations.extend(var_list)
                    # Add the original skill
                    variations.append(skill_lower)
                    # Remove duplicates
                    variations = list(set(variations))
                    
                    for var in variations:
                        skill_texts.append(var)
                        skill_map[var] = skill
                
                # Convert to embeddings
                skill_embeddings = self.embedding_model.encode(skill_texts)
                resume_embeddings = self.embedding_model.encode(resume_sentences)
                
                # Calculate similarities
                similarities = resume_embeddings @ skill_embeddings.T
                max_similarities = np.max(similarities, axis=0)
                
                # Use a moderate threshold for semantic matching
                threshold = 0.6  # Higher threshold for more precise matching
                
                matched_variations = set()
                for idx, skill_text in enumerate(skill_texts):
                    if max_similarities[idx] > threshold:
                        original_skill = skill_map[skill_text]
                        if original_skill not in matched_skills:
                            matched_skills.append(original_skill)
                            matched_variations.add(skill_text)
            
            # Get missing skills
            missing_skills = [skill for skill in job_requirements['skills'] if skill not in matched_skills]
            
            return {
                'matched_skills': matched_skills,
                'missing_skills': missing_skills
            }
        except Exception as e:
            st.error(f"Error in skill matching: {str(e)}")
            return {
                'matched_skills': [],
                'missing_skills': []
            }
    
    def tailor_resume(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> Dict:
        """Generate a tailored resume using the LLM and provide improvement analysis."""
        resume_prompt = f"""You are an expert ATS optimization specialist. Rewrite the following resume to maximize its ATS score while maintaining readability.
        The goal is to significantly improve the resume's ATS score by incorporating job-specific keywords and requirements.

        Job Requirements to Target:
        1. Required Skills: {job_requirements.get('skills', [])}
        2. Experience Needed: {job_requirements.get('experience', 'Not specified')}
        3. Responsibilities: {job_requirements.get('responsibilities', [])}

        Current Status:
        - Matched Skills: {skill_matches['matched_skills']}
        - Missing Skills: {skill_matches['missing_skills']}

        Optimization Requirements:
        1. Keyword Integration:
           - Add ALL missing required skills with relevant context
           - Place important keywords in prominent positions
           - Use exact phrases from job requirements
           - Maintain optimal keyword density (5-8%)
        
        2. Format Optimization:
           - Use clear section headers: Summary, Experience, Skills, Education
           - Start bullets with strong action verbs
           - Ensure consistent formatting
           - Use standard bullet points
        
        3. Content Enhancement:
           - Add quantifiable metrics to achievements
           - Highlight experience matching job requirements
           - Emphasize transferable skills
           - Use industry-standard terminology
        
        4. ATS Guidelines:
           - Use full terms before abbreviations
           - Avoid tables, columns, and graphics
           - Use standard job titles
           - Place keywords near the start of bullet points

        Original Resume:
        {resume_text}
        
        Job Requirements:
        {json.dumps(job_requirements, indent=2)}
        
        Return ONLY the optimized resume text. Ensure EVERY required skill and responsibility is addressed.
        """
        
        tailored_resume = str(self.llm.invoke(resume_prompt).content)
        
        # Then, get the analysis separately
        analysis_prompt = f"""Analyze how the resume matches the job requirements and provide a detailed improvement analysis.
        
        IMPORTANT RULES:
        1. ONLY mention skills that are EXPLICITLY stated in the resume
        2. DO NOT make assumptions about skills not directly mentioned
        3. DO NOT infer skills from project descriptions unless explicitly stated
        4. If a skill is missing, list it in missing skills, do not try to find similar alternatives
        5. For matched skills, quote the exact text from resume that demonstrates the skill
        
        Focus on these aspects:
        1. Skills alignment - EXACT matches only
        2. Experience relevance - DIRECT matches only
        3. Achievement emphasis - ACTUAL achievements mentioned
        4. Missing keywords - List ALL required skills not found in resume
        
        Original Resume:
        {resume_text}
        
        Job Requirements:
        {json.dumps(job_requirements, indent=2)}
        
        Currently Matched Skills (verified): {skill_matches['matched_skills']}
        Currently Missing Skills (verified): {skill_matches['missing_skills']}
        
        Return a JSON object with this exact structure:
        {{
            "improvements": [
                "specific improvements needed based on ACTUAL gaps"
            ],
            "skills_analysis": {{
                "matched": [
                    "ONLY skills explicitly found in resume with exact quotes"
                ],
                "missing": [
                    "ONLY skills from job requirements that are completely absent from resume"
                ]
            }},
            "achievement_emphasis": [
                "ONLY quantifiable achievements actually present in resume"
            ],
            "keyword_optimization": [
                "ONLY keywords from job requirements that should be added"
            ]
        }}
        
        Remember: Do not infer, assume, or suggest skills that are not explicitly stated in the resume.
        """
        
        analysis_response = self.llm.invoke(analysis_prompt)
        try:
            content = str(analysis_response.content).strip()
            # More robust JSON extraction
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]
            content = content.strip()
            
            try:
                analysis_result = json.loads(content)
                analysis_result["tailored_resume"] = tailored_resume
                return analysis_result
            except json.JSONDecodeError as json_err:
                return {
                    "improvements": ["Error analyzing improvements"],
                    "skills_analysis": {
                        "matched": ["Error analyzing matched skills"],
                        "missing": ["Error analyzing missing skills"]
                    },
                    "achievement_emphasis": ["Error analyzing achievements"],
                    "keyword_optimization": ["Error analyzing keywords"],
                    "tailored_resume": tailored_resume
                }
        except Exception as e:
            return {
                "improvements": ["Error analyzing improvements"],
                "skills_analysis": {
                    "matched": ["Error analyzing matched skills"],
                    "missing": ["Error analyzing missing skills"]
                },
                "achievement_emphasis": ["Error analyzing achievements"],
                "keyword_optimization": ["Error analyzing keywords"],
                "tailored_resume": tailored_resume
            }
    
    def generate_docx(self, tailored_content: str) -> Document:
        """Convert the tailored content into a DOCX file."""
        doc = Document()
        # Only include the actual resume content, no analysis
        content = str(tailored_content).strip()
        for paragraph in content.split('\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        return doc
    
    def generate_cold_email(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> str:
        """Generate a personalized cold email based on the resume and job requirements."""
        prompt = f"""Generate a professional cold email for a job application based on the following information.
        The email should:
        1. Be concise and professional
        2. Highlight key matching skills: {skill_matches['matched_skills']}
        3. Reference specific job requirements
        4. Include a brief mention of relevant achievements
        5. End with a call to action
        
        Job Details:
        Company: {job_requirements.get('company', 'the company')}
        Position: {job_requirements.get('title', 'the position')}
        
        Resume Context:
        {resume_text}  # use resume_text for context
        
        Format the email with:
        - Professional subject line
        - Greeting
        - 2-3 concise paragraphs
        - Professional closing
        - keep it concise
        Return the complete email with subject line.
        """
        
        response = self.llm.invoke(prompt)
        return str(response.content)

    def calculate_ats_score(self, resume_text: str, job_requirements: Dict, skill_matches: Dict) -> Dict:
        """Calculate a comprehensive ATS score for the resume based on multiple factors."""
        prompt = f"""You are an ATS (Applicant Tracking System) expert. Analyze the resume against the job requirements and calculate scores.

        Rules for scoring:
        1. All scores must be integers (whole numbers)
        2. Each section score must not exceed its maximum value
        3. Total score must be the sum of all section scores
        4. Compare directly against job requirements
        5. Higher scores for exact keyword matches from requirements

        Scoring Criteria:
        1. Keyword Match (30 points max):
           - Award points for exact matches with job requirements
           - Check keyword frequency and placement
           - Required skills present: {job_requirements.get('skills', [])}
           - Current matched skills: {skill_matches['matched_skills']}

        2. Experience Alignment (25 points max):
           - Compare against required experience: {job_requirements.get('experience', 'Not specified')}
           - Check for relevant role titles
           - Evaluate described responsibilities against: {job_requirements.get('responsibilities', [])}

        3. Skills Match (25 points max):
           - Technical skills alignment
           - Soft skills presence
           - Skills context and application

        4. Education Relevance (10 points max):
           - Required education level match
           - Field of study relevance
           - Certifications value

        5. Format & Organization (10 points max):
           - Standard section headers
           - Bullet point structure
           - Content readability

        Resume to analyze:
        {resume_text}

        Job Requirements:
        {json.dumps(job_requirements, indent=2)}

        Return a JSON object with this exact structure:
        {{
            "total_score": <integer 0-100>,
            "section_scores": {{
                "keyword_match": {{
                    "score": <integer 0-30>,
                    "max": 30,
                    "details": ["<specific keywords found>", "<specific keywords missing>"]
                }},
                "experience": {{
                    "score": <integer 0-25>,
                    "max": 25,
                    "details": ["<specific experience matches>", "<experience gaps>"]
                }},
                "skills": {{
                    "score": <integer 0-25>,
                    "max": 25,
                    "details": ["<matched skills details>", "<missing skills impact>"]
                }},
                "education": {{
                    "score": <integer 0-10>,
                    "max": 10,
                    "details": ["<education alignment details>"]
                }},
                "format": {{
                    "score": <integer 0-10>,
                    "max": 10,
                    "details": ["<format strengths>", "<format improvements needed>"]
                }}
            }},
            "improvement_suggestions": [
                "<actionable suggestion 1>",
                "<actionable suggestion 2>",
                "<actionable suggestion 3>"
            ],
            "keyword_density": {{
                "<actual keyword from job requirements>": <integer frequency>
            }}
        }}
        """

        try:
            response = self.llm.invoke(prompt)
            content = str(response.content).strip()
            
            # More robust JSON extraction
            if "```json" in content:
                content = content.split("```json")[1].split("```")[0]
            elif "```" in content:
                content = content.split("```")[1].split("```")[0]
            content = content.strip()
            
            # Parse and validate the response
            result = json.loads(content)
            
            # Validate scores are integers and within range
            result["total_score"] = int(result["total_score"])
            for section, data in result["section_scores"].items():
                data["score"] = int(data["score"])
                if data["score"] < 0 or data["score"] > data["max"]:
                    raise ValueError(f"Invalid score for {section}: {data['score']}")
            
            # Validate total score matches sum of sections
            expected_total = sum(data["score"] for data in result["section_scores"].values())
            if result["total_score"] != expected_total:
                result["total_score"] = expected_total
            
            return result
            
        except (json.JSONDecodeError, KeyError, ValueError) as e:
            st.error(f"Error in ATS scoring: {str(e)}")
            return {
                "total_score": 0,
                "section_scores": {
                    "keyword_match": {"score": 0, "max": 30, "details": ["Error analyzing keywords"]},
                    "experience": {"score": 0, "max": 25, "details": ["Error analyzing experience"]},
                    "skills": {"score": 0, "max": 25, "details": ["Error analyzing skills"]},
                    "education": {"score": 0, "max": 10, "details": ["Error analyzing education"]},
                    "format": {"score": 0, "max": 10, "details": ["Error analyzing format"]}
                },
                "improvement_suggestions": ["Unable to generate suggestions due to an error"],
                "keyword_density": {}
            }

def main():
    st.set_page_config(layout="wide")  # Use wide layout for full screen
    
    # Center-aligned title with custom styling
    st.markdown("""
        <h1 style='text-align: center; color: #1E88E5; padding: 1rem;'>
            🎯 AI-Powered Resume Tailor
        </h1>
        <p style='text-align: center; color: #666; margin-bottom: 2rem;'>
            Optimize your resume for ATS using AI
        </p>
    """, unsafe_allow_html=True)
    
    try:
        # Initialize ResumeTailor
        tailor = ResumeTailor()
        
        # Create two main columns for the layout
        left_col, right_col = st.columns([1, 1.5], gap="large")  # Added gap between columns
        
        with left_col:
            st.header("📤 Upload Resume & Job Details")
            # File upload
            resume_file = st.file_uploader("Upload your resume (PDF/DOCX)", type=['pdf', 'docx'])
            
            # Job input method selection
            job_input_method = st.radio(
                "Choose how to input the job description:",
                ["Enter URL", "Paste Description"]
            )
            
            job_text = ""
            is_url = True
            
            if job_input_method == "Enter URL":
                job_text = st.text_input("Enter the job posting URL:")
                is_url = True
            else:
                job_text = st.text_area("Paste the job description here:", height=300)
                is_url = False
            
            if resume_file and job_text:
                process_button = st.button("🎯 Tailor Resume", use_container_width=True)
            else:
                st.info("Please upload your resume and provide job details to proceed.")
                process_button = False
        
        with right_col:
            if resume_file and job_text and process_button:
                with st.spinner("Processing your resume..."):
                    # Extract resume text
                    if resume_file.type == "application/pdf":
                        resume_text = tailor.extract_text_from_pdf(resume_file)
                    else:
                        resume_text = tailor.extract_text_from_docx(resume_file)
                    
                    # Parse job description
                    job_requirements = tailor.parse_job_description(job_text, is_url)
                    
                    # Match skills
                    skill_matches = tailor.match_skills(resume_text, job_requirements)
                    
                    # Calculate initial ATS score (before tailoring)
                    initial_ats_score = tailor.calculate_ats_score(resume_text, job_requirements, skill_matches)
                    
                    # Generate tailored resume with analysis
                    analysis_result = tailor.tailor_resume(resume_text, job_requirements, skill_matches)
                    
                    # Calculate final ATS score (after tailoring)
                    final_ats_score = tailor.calculate_ats_score(analysis_result['tailored_resume'], job_requirements, skill_matches)
                    
                    # Generate cold email
                    cold_email = tailor.generate_cold_email(resume_text, job_requirements, skill_matches)
                    
                    # Create tabs for different sections
                    tabs = st.tabs(["💡 Analysis", "📊 ATS Score", "📧 Cold Email", "📄 Resume Versions"])
                    
                    # Analysis Tab
                    with tabs[0]:
                        # Skills Analysis
                        st.subheader("🎯 Skills Analysis")
                        skill_cols = st.columns(2)
                        with skill_cols[0]:
                            st.write("✅ Matched Skills")
                            for idx, skill_analysis in enumerate(analysis_result['skills_analysis']['matched'], 1):
                                st.write(f"{idx}. {skill_analysis}")
                        with skill_cols[1]:
                            st.write("📝 Missing Skills & Suggestions")
                            for idx, skill_suggestion in enumerate(analysis_result['skills_analysis']['missing'], 1):
                                st.write(f"{idx}. {skill_suggestion}")
                        
                        # Improvements
                        st.subheader("🔄 Improvements Made")
                        for idx, improvement in enumerate(analysis_result['improvements'], 1):
                            st.write(f"{idx}. {improvement}")
                        
                        # Achievements
                        st.subheader("📊 Quantifiable Achievements")
                        for idx, achievement in enumerate(analysis_result['achievement_emphasis'], 1):
                            st.write(f"{idx}. {achievement}")
                        
                        # Keywords
                        st.subheader("🔍 ATS Keywords")
                        for idx, keyword in enumerate(analysis_result['keyword_optimization'], 1):
                            st.write(f"{idx}. {keyword}")
                    
                    # ATS Score Tab
                    with tabs[1]:
                        st.subheader("📊 ATS Score Analysis")
                        
                        # Display before and after scores side by side
                        score_cols = st.columns(2)
                        
                        # Before Score
                        with score_cols[0]:
                            st.markdown("### Before Tailoring")
                            initial_total = initial_ats_score["total_score"]
                            st.markdown(f"""
                                <div style='text-align: center;'>
                                    <h1 style='color: {"#28a745" if initial_total >= 70 else "#ffc107" if initial_total >= 50 else "#dc3545"}; font-size: 4rem;'>
                                        {initial_total}/100
                                    </h1>
                                </div>
                            """, unsafe_allow_html=True)
                        
                        # After Score
                        with score_cols[1]:
                            st.markdown("### After Tailoring")
                            final_total = final_ats_score["total_score"]
                            st.markdown(f"""
                                <div style='text-align: center;'>
                                    <h1 style='color: {"#28a745" if final_total >= 70 else "#ffc107" if final_total >= 50 else "#dc3545"}; font-size: 4rem;'>
                                        {final_total}/100
                                    </h1>
                                </div>
                            """, unsafe_allow_html=True)
                        
                        # Display improvement percentage
                        if initial_total > 0:  # Avoid division by zero
                            improvement = ((final_total - initial_total) / initial_total) * 100
                            st.markdown(f"""
                                <div style='text-align: center; margin: 20px 0;'>
                                    <h3 style='color: {"#28a745" if improvement > 0 else "#dc3545"}'>
                                        {'+' if improvement > 0 else ''}{improvement:.1f}% Improvement
                                    </h3>
                                </div>
                            """, unsafe_allow_html=True)
                        
                        # Display section scores comparison
                        st.subheader("Section Scores Comparison")
                        for section, data in final_ats_score["section_scores"].items():
                            st.write(f"**{section.replace('_', ' ').title()}**")
                            cols = st.columns([3, 3, 1])
                            
                            # Before score
                            with cols[0]:
                                initial_progress = initial_ats_score["section_scores"][section]["score"] / data["max"]
                                st.progress(initial_progress, text=f"Before: {initial_ats_score['section_scores'][section]['score']}/{data['max']}")
                            
                            # After score
                            with cols[1]:
                                final_progress = data["score"] / data["max"]
                                st.progress(final_progress, text=f"After: {data['score']}/{data['max']}")
                            
                            # Details button
                            with cols[2]:
                                if st.button(f"Details 🔍", key=f"details_{section}"):
                                    st.write("Before:", initial_ats_score["section_scores"][section]["details"])
                                    st.write("After:", data["details"])
                        
                        # Display keyword density
                        st.subheader("🔑 Keyword Density Comparison")
                        density_cols = st.columns(2)
                        
                        with density_cols[0]:
                            st.write("Before Tailoring")
                            if initial_ats_score["keyword_density"]:
                                initial_keywords_df = pd.DataFrame(
                                    list(initial_ats_score["keyword_density"].items()),
                                    columns=["Keyword", "Frequency"]
                                ).sort_values(by="Frequency", ascending=False)
                                st.dataframe(initial_keywords_df, use_container_width=True)
                        
                        with density_cols[1]:
                            st.write("After Tailoring")
                            if final_ats_score["keyword_density"]:
                                final_keywords_df = pd.DataFrame(
                                    list(final_ats_score["keyword_density"].items()),
                                    columns=["Keyword", "Frequency"]
                                ).sort_values(by="Frequency", ascending=False)
                                st.dataframe(final_keywords_df, use_container_width=True)
                        
                        # Display improvement suggestions
                        st.subheader("📈 Improvement Suggestions")
                        for idx, suggestion in enumerate(final_ats_score["improvement_suggestions"], 1):
                            st.write(f"{idx}. {suggestion}")
                    
                    # Cold Email Tab
                    with tabs[2]:
                        st.subheader("📧 Cold Email Template")
                        st.text_area("Copy the email below:", cold_email, height=400)
                    
                    # Resume Versions Tab
                    with tabs[3]:
                        st.subheader("📋 Compare Versions")
                        version_cols = st.columns(2)
                        with version_cols[0]:
                            st.write("Original Resume")
                            st.text_area("", resume_text, height=400, disabled=True)
                        with version_cols[1]:
                            st.write("Tailored Resume")
                            st.text_area("", analysis_result['tailored_resume'], height=400, disabled=True)
                        
                        # Download button
                        st.subheader("📥 Download")
                        doc = tailor.generate_docx(analysis_result['tailored_resume'])
                        doc.save("tailored_resume.docx")
                        
                        with open("tailored_resume.docx", "rb") as file:
                            st.download_button(
                                label="⬇️ Download Tailored Resume",
                                data=file,
                                file_name="tailored_resume.docx",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                use_container_width=True
                            )
                
    except ValueError as e:
        st.error(f"Configuration Error: {str(e)}")
        st.info("Please ensure you have set up the GROQ_API_KEY in your .env file")

if __name__ == "__main__":
    main()
